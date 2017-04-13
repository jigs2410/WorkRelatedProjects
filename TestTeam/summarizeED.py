# -*- coding: utf-8 -*-
"""
Created on Tue Apr 11 11:40:29 2017

@author: jrathod
"""
## do this pip install --pre python-docx
import re
from docx import Document
import collections
from collections import OrderedDict

TOdict=OrderedDict()

wordDoc = Document('sample.docx')
TC=""
print(len(wordDoc.tables))

for table in wordDoc.tables:
    rowCount=0
    for row in table.rows:
        if rowCount ==0:
            cells = row.cells
            TC = cells[0].text.strip(' \t\n\r')
            print(cells[0].text.strip(' \t\n\r'))
        if rowCount >1:
        #if row[0].cell[0].text!="Step No." and row[1]!="Test steps : Actions":
            cells = row.cells
            step = re.sub('\s+',' ', cells[0].text.strip(' \t\n\r')) 
            to_id = re.sub('\s+',' ', cells[3].text.strip(' \t\n\r'))
            to_id = re.sub('[A-Za-z-]+','',to_id)
            #to_id = to_id.strip(' \t\n\r')
            to_id = to_id.split(' ')
            to_id = list(filter(None, to_id))
            # filter does this-> (item for item in iterable if function(item)) if function is not None and (item for item in iterable if item) if function is None.
            rule_ids = re.sub('\s+',' ', cells[5].text.strip(' \t\n\r'))
            rule_ids = rule_ids.split(' ') ## make sure ED's are unique
            ruleset = set(rule_ids)
            rule_ids = list(ruleset)
                    
            #print(cells[0].text.strip(' \t\n\r'),cells[3].text.strip(' \t\n\r'),cells[4].text.strip(' \t\n\r'))
            
            if len(to_id) > 0:
                print("\t",step,to_id,rule_ids)
                for id in to_id:
                    print("Processing %s right now "%id)
                    if id in TOdict:
                        print("present")
                        print("id is ->", id ,TOdict[id] )
                        rule_idsSoFar = TOdict[id]
                        rule_idsSoFar = rule_idsSoFar + rule_ids
                        rule_idsSoFarset = set(rule_idsSoFar)
                        rule_idsSoFar = list(rule_idsSoFarset)
                        TOdict[id] = rule_idsSoFar
                    else: #insert
                        TOdict[id] = rule_ids
                #print(to_id)
                #print(rule_ids)
                
        rowCount+=1
        
for key,value in TOdict.items():
    value = sorted(value)
    print(key,"->",value)
    