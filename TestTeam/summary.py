# -*- coding: utf-8 -*-
"""
Created on Wed Apr 12 14:18:04 2017

@author: jrathod
"""
## do this pip install --pre python-docx
## if want to implement lot of records with memory consumption less than 10 MB use write only mode
import re
from docx import Document
import collections
from collections import OrderedDict

from openpyxl import Workbook
from openpyxl.compat import range
from openpyxl.utils import get_column_letter
from openpyxl import load_workbook

TOdict=OrderedDict()

SYRS_TO = OrderedDict()
UCName=""
TestCase=""
wb = Workbook()
ws1 = wb.active
ws1.title = "Traceability"
ws1.cell(column=1, row=1, value="SYRS NUMBER")
ws1.cell(column=2, row=1, value="SYRSTO-ID")
ws1.cell(column=3, row=1, value="DATA DICTIONARY RULE ID")
ws1.cell(column=4, row=1, value="TEST CASE NUMBER") 
#ws1.cell(column=5, row=1, value="PUMPS")
validDoc=0

excelRow = 2
 
# write to open excel file
def writeToExcel(TO,EDList,TestCase):
    global ws1
    global excelRow
    currentRow = excelRow
    if TO in SYRS_TO:
        #print(SYRS_TO[TO],TO,EDList,TestCase)
        to_ids = SYRS_TO[TO]
        
        # to_ids are more than one write it on different lines
        if type(to_ids) is list: 
            #print("list", to_ids)
            for id in to_ids:
                ws1.cell(column=1, row=currentRow, value=id)
                currentRow+=1
        else:
            ws1.cell(column=1, row=excelRow, value=to_ids)
            
        ws1.cell(column=2, row=excelRow, value=TO)
        ws1.cell(column=4, row=excelRow, value=TestCase)
        
        # ED_ID's are more than one write it on different lines
        currentRow = excelRow
        for id in EDList:
            ws1.cell(column=3, row=currentRow, value=id)
            currentRow+=1
            
        #preserve the rowcounter
        if excelRow<currentRow:
            excelRow=currentRow
        else:    
            excelRow+=1
    else:
        print("TO %s is not defined in Objective Table"%TO)
        
testCaseDoc = input("Enter TestCase document path :- ")
targetFileDir=input("Enter where to store traceability excel file :- ")
targetFile=""
#wordDoc = Document(r'C:\Users\jrathod\Desktop\WorkDocs\TestTeam\Template1-TestCases-Traceability.docx')
try:
    wordDoc = Document(testCaseDoc)
except:
    print("Test case doc not found at path : ",testCaseDoc)
else:    
    
    
#print(len(wordDoc.tables))
## check for tables if table count includes usecase name , test object and 1 test case.
## so min count is 3

    tableCount = len(wordDoc.tables)
    tableCounter = 1
    if tableCount  < 3:
        print (" Invalid Format make sure UseCase , Test Objectives and Test Case are present in doc")
    else: # valid document start processing
        #print("Good doc")
        for table in wordDoc.tables:
            if tableCounter == 1: #usecase table
                 rowCount=0
                 for row in table.rows:
                    if rowCount ==0:
                        cells = row.cells
                        if (cells[0].text.strip(' \t\n\r') == "Name of Use Case"):
                            UCName = cells[1].text.strip(' \t\n\r')
                            print("UseCase Name is ",UCName)
                            if (targetFileDir[:len(targetFileDir)]) == "\\":
                                targetFile = targetFileDir+"Traceability_"+UCName+".xlsx"
                            else:
                                targetFile = targetFileDir+"\\Traceability_"+UCName+".xlsx"
                            validDoc = 1 #if doc is valid till now 
                        else:
                            print("Name of Use Case is not defined in correct format")
                            break
            
            if tableCounter == 2 and validDoc == 1: #TestObjective Table
                #print("now processing TO table")
                rowCount=0
                for row in table.rows:
                    cells = row.cells
                    col1 = re.sub('\s+',' ', cells[0].text.strip(' \t\n\r')) 
                    col2 = re.sub('\s+',' ', cells[1].text.strip(' \t\n\r')) 
                    col3 = re.sub('\s+',' ', cells[2].text.strip(' \t\n\r'))
                    if rowCount ==0:
                        if col1 =="SYRSTO ID Number" and col2=="Test Objective (TO)" and col3=="Requirements Number (SyRS #)":
                             #print("col1",col1,"col2",col2,"col3",col3)
                             validDoc= 2 # increment validDoc counter
                        else:
                            print("Test Objectives table missing or header is not correct")
                            break
                    else:
                        #print("processing other rows")
                        SYRS_List=[]
                        if (len(col1)>0): # if data is present process else ignore
                            TOObj = re.match( r'(\d*).*', col1, re.I)
                            if TOObj:
                                TO_ID =TOObj.group(1)
                                SYRSObj = re.match( r'(\d*).*', col3, re.I)
                                if SYRSObj:
                                    SYRS =SYRSObj.group(1)
                                    if (len(TO_ID)>0 and len(SYRS)>0):
                                        #print("TO_ID",TO_ID ,"SYRS", SYRS )
                                        if TO_ID in SYRS_TO:
                                            #print("TO_ID already present",TO_ID)
                                            SYRS_List.append(SYRS_TO[TO_ID])
                                            SYRS_List.append(SYRS)
                                            SYRS_TO[TO_ID] = SYRS_List
                                        else:
                                            SYRS_TO[TO_ID] = SYRS
                                
                    rowCount+=1      
                          
            
            if tableCounter >= 3 and validDoc == 2: #Process Test Case
                TestCase=""
                #print("now processing TO table")
                rowCount=0
                #print("length of table is " , len(table.rows))
                for row in table.rows:
                    cells = row.cells
                    if rowCount ==0:
                        if "Test Case" in cells[0].text.strip(' \t\n\r'):
                            TestCase = cells[0].text.strip(' \t\n\r')
                            print("Processing TestCase:-",TestCase)
                        else:
                            print("Invalid Test case table format. Missing Test Case definition")
                            validDoc = 0 # reset so we dont process any table until this is corrected
                            break
                            
                         #print(cells[0].text.strip(' \t\n\r'))
                         #print("Processing Test Case ->",TestCase)
                    
                    if rowCount ==1:
                        stepNoHeader = cells[0].text.strip(' \t\n\r')
                        syrsTOHeader = cells[3].text.strip(' \t\n\r')
                        ruleIDHeader = cells[4].text.strip(' \t\n\r')
                        #print(stepNoHeader,syrsTOHeader,ruleIDHeader)
                        if not ("Step No" in stepNoHeader) or not("SYRSTO" in syrsTOHeader) or not("Rule Id" in ruleIDHeader ):
                            print("Invalid Test case table format. Missing Header definition")
                            validDoc = 0 # reset so we dont process any table until this is corrected
                            break
                    #process the data in test case and add to our TO dict   
                    if rowCount > 1:
                        to_id = re.sub('\s+',' ', cells[3].text.strip(' \t\n\r'))
                        to_id = re.sub('[A-Za-z-]+','',to_id)
                        #to_id = to_id.strip(' \t\n\r')
                        to_id = to_id.split(' ')
                        to_id = list(filter(None, to_id))
                        # filter does this-> (item for item in iterable if function(item)) if function is not None and (item for item in iterable if item) if function is None.
                        rule_ids = re.sub('\s+',' ', cells[4].text.strip(' \t\n\r'))
                        if len(rule_ids) > 0: #ignore empty rules
                            rule_ids = rule_ids.split(' ') ## make sure ED's are unique
                            ruleset = set(rule_ids)
                            rule_ids = list(ruleset)
                        
                        #print(to_id,rule_ids)
                        #make sure to_id and rule_ids are present
                        if len(to_id) > 0 and len(rule_ids) > 0:
                            #print("\t",to_id,rule_ids,"->",len(rule_ids))
                            for id in to_id:
                                #print("Processing %s right now "%id)
                                if id in TOdict:
                                    #print("id is ->", id ,TOdict[id] )
                                    rule_idsSoFar = TOdict[id]
                                    rule_idsSoFar = rule_idsSoFar + rule_ids
                                    rule_idsSoFarset = set(rule_idsSoFar)
                                    rule_idsSoFar = list(rule_idsSoFarset)
                                    TOdict[id] = rule_idsSoFar
                                else: #insert
                                    TOdict[id] = rule_ids
                                
                    rowCount+=1
                    # at the end of testcase table enter it in TOid
                    if rowCount == len(table.rows):
                        #print("Counter",rowCount,"length",len(table.rows))
                        keyList = TOdict.keys()
                        keyList = sorted(keyList)
                        for key in keyList:
                            value = TOdict[key]
                            value = sorted(value)
                        #for key,value in TOdict.items():
                            #value = sorted(value)
                            print("\t",key,"->",value)
                            writeToExcel(key,value,TestCase)
                            
                        TOdict.clear()
            tableCounter+=1 # increment table counter
        
    #    for key,value in SYRS_TO.items():
    #        #value = sorted(value)
    #        print(key,"->",value)
            
        if validDoc >= 1:
            wb.save(filename = targetFile)